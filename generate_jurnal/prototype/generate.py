import spacy
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

nlp = spacy.load("en_core_web_sm")

def baca_template(path_template):
    dokumen_template = Document(path_template)
    return dokumen_template

def ekstrak_judul(dokumen):
    for para in dokumen.paragraphs:
        if para.text.strip():
            return para.text.strip()
    return "Judul Tidak Ditemukan"

def ekstrak_bagian_dengan_nlp(dokumen):
    bagian = {
        "Judul": "",
        "Abstrak": [],
        "Abstract": [],
        "Pendahuluan": [],
        "Metode Penelitian": [],
        "Hasil dan Pembahasan": [],
        "Kesimpulan": [],
        "Referensi": [],
        "Keywords": "",
        "Kata Kunci": "",
    }

    bagian["Judul"] = ekstrak_judul(dokumen)
    
    section_flags = {
        "dalam_abstrak": False,
        "dalam_abstract": False,
        "dalam_pendahuluan": False,
        "dalam_metode_penelitian": False,
        "dalam_hasil_dan_pembahasan": False,
        "dalam_kesimpulan": False,
        "dalam_referensi": False
    }

    current_section = None
    current_list_pendahuluan = []
    current_list_metode_penelitian = []
    current_list_hasil_dan_pembahasan = []
    current_list_kesimpulan = []
    current_list_referensi = []

    def reset_flags(flags):
        for key in flags:
            flags[key] = False

    for para in dokumen.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if any(keyword in text for keyword in ["Keywords", "Keyword", "Kata Kunci", "UCAPAN TERIMA KASIH", "BAB", "LAMPIRAN"]):
            reset_flags(section_flags)
            current_section = None
            continue
        
        if "ABSTRACT" in text:
            current_section = "Abstract"
            reset_flags(section_flags)
            section_flags["dalam_abstract"] = True
        elif "ABSTRAK" in text:
            current_section = "Abstrak"
            reset_flags(section_flags)
            section_flags["dalam_abstrak"] = True        
        elif "PENDAHULUAN" in text:
            current_section = "Pendahuluan"
            reset_flags(section_flags)
            section_flags["dalam_pendahuluan"] = True
        elif "METODE PENELITIAN" in text:
            current_section = "Metode Penelitian"
            reset_flags(section_flags)
            section_flags["dalam_metode_penelitian"] = True
        elif "HASIL DAN DISKUSI" in text or "HASIL DAN PEMBAHASAN" in text:
            current_section = "Hasil dan Pembahasan"
            reset_flags(section_flags)
            section_flags["dalam_hasil_dan_pembahasan"] = True
        elif "KESIMPULAN" in text or "PENUTUP" in text:
            current_section = "Kesimpulan"
            reset_flags(section_flags)
            section_flags["dalam_kesimpulan"] = True
        elif "DAFTAR PUSTAKA" in text or "REFERENSI" in text:
            current_section = "Referensi"
            reset_flags(section_flags)
            section_flags["dalam_referensi"] = True
        elif para.style.name.startswith("Heading"):
            continue
        else:
            if section_flags["dalam_pendahuluan"]:
                if para.style.name.startswith("List"):
                    current_list_pendahuluan.append(text)
                else:
                    if current_list_pendahuluan:
                        bagian["Pendahuluan"].append(current_list_pendahuluan)
                        current_list_pendahuluan = []
                    bagian["Pendahuluan"].append(text)
            elif section_flags["dalam_metode_penelitian"]:
                if para.style.name.startswith("List"):
                    current_list_metode_penelitian.append(text)
                else:
                    if current_list_metode_penelitian:
                        bagian["Metode Penelitian"].append(current_list_metode_penelitian)
                        current_list_metode_penelitian = []
                    bagian["Metode Penelitian"].append(text)
            elif section_flags["dalam_hasil_dan_pembahasan"]:
                if para.style.name.startswith("List"):
                    current_list_hasil_dan_pembahasan.append(text)
                else:
                    if current_list_hasil_dan_pembahasan:
                        bagian["Hasil dan Pembahasan"].append(current_list_hasil_dan_pembahasan)
                        current_list_hasil_dan_pembahasan = []
                    bagian["Hasil dan Pembahasan"].append(text)
            elif section_flags["dalam_kesimpulan"]:
                if para.style.name.startswith("List"):
                    current_list_kesimpulan.append(text)
                else:
                    if current_list_kesimpulan:
                        bagian["Kesimpulan"].append(current_list_kesimpulan)
                        current_list_kesimpulan = []
                    bagian["Kesimpulan"].append(text)
            elif section_flags["dalam_referensi"]:
                if para.style.name.startswith("List"):
                    current_list_referensi.append(text)
                else:
                    if current_list_referensi:
                        bagian["Referensi"].append(current_list_referensi)
                        current_list_referensi = []
                    bagian["Referensi"].append(text)
            else:
                if current_section:
                    bagian[current_section].append(text)

    if current_list_pendahuluan:
        bagian["Pendahuluan"].append(current_list_pendahuluan)
    if current_list_metode_penelitian:
        bagian["Metode Penelitian"].append(current_list_metode_penelitian)
    if current_list_hasil_dan_pembahasan:
        bagian["Hasil dan Pembahasan"].append(current_list_hasil_dan_pembahasan)
    if current_list_kesimpulan:
        bagian["Kesimpulan"].append(current_list_kesimpulan)
    if current_list_referensi:
        bagian["Referensi"].append(current_list_referensi)

    teks_dokumen = "\n".join([para.text for para in dokumen.paragraphs])
    nlp_dokumen = nlp(teks_dokumen)
    kata_kunci = [token.text for token in nlp_dokumen if token.pos_ in ["NOUN", "PROPN"]]
    bagian["Kata Kunci"] = "; ".join(set(kata_kunci[:5])) 

    return bagian

def sesuaikan_dengan_template(dokumen_template, bagian):
    for para in dokumen_template.paragraphs:
        if "JUD1" in para.text:
            para.clear()
            judul = bagian.get('Judul', 'Judul Tidak Ditemukan').upper() 
            if len(judul.split()) > 15: 
                judul = ' '.join(judul.split()[:15])
            p_judul = para.insert_paragraph_before(judul)
            run_judul = p_judul.runs[0]
            run_judul.font.size = Pt(14)
            run_judul.font.name = 'Arial'
            run_judul.bold = True
            p_judul.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        elif "AA1" in para.text:
            para.clear()
            for paragraph in bagian.get('Abstract', ['Tidak ada abstrak ditemukan.']):
                p = para.insert_paragraph_before(paragraph)
                run = p.runs[0]
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                run.italic = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.first_line_indent = Inches(0.5)

        elif "AA2" in para.text:
            para.clear()
            for paragraph in bagian.get('Abstrak', ['Tidak ada abstrak ditemukan.']):
                p = para.insert_paragraph_before(paragraph)
                p.style.font.size = Pt(10)
                p.style.font.name = 'Arial'
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.first_line_indent = Inches(0.5)
        
        elif "AA3" in para.text:
            para.clear()
            for item_pendahuluan in bagian.get('Pendahuluan', ['Tidak ada pendahuluan ditemukan.']):
                if isinstance(item_pendahuluan, list):
                    for i, list_item in enumerate(item_pendahuluan, 1):
                        p1 = para.insert_paragraph_before(f"{i}. {list_item}")
                        p1.style.font.size = Pt(10)
                        p1.style.font.name = 'Arial'
                        p1.paragraph_format.left_indent = Inches(0.5)
                        p1.paragraph_format.first_line_indent = Inches(-0.15)
                        p1.paragraph_format.space_before = Pt(3)
                        p1.paragraph_format.space_after = Pt(3)
                        p1.paragraph_format.line_spacing = 1.15
                        p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p1 = para.insert_paragraph_before(item_pendahuluan)
                    p1.style.font.size = Pt(10)
                    p1.style.font.name = 'Arial'
                    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p1.paragraph_format.space_before = Pt(3)
                    p1.paragraph_format.space_after = Pt(3)
                    p1.paragraph_format.line_spacing = 1.15
                    p1.paragraph_format.first_line_indent = Inches(0.5)

        elif "AA4" in para.text:
            para.clear()
            for item_metode_penelitian in bagian.get('Metode Penelitian', ['Tidak ada metode penelitian ditemukan.']):
                if isinstance(item_metode_penelitian, list):
                    for i, list_item in enumerate(item_metode_penelitian, 1):
                        p2 = para.insert_paragraph_before(f"{i}. {list_item}")
                        p2.style.font.size = Pt(10)
                        p2.style.font.name = 'Arial'
                        p2.paragraph_format.left_indent = Inches(0.5)
                        p2.paragraph_format.first_line_indent = Inches(-0.15)
                        p2.paragraph_format.space_before = Pt(3)
                        p2.paragraph_format.space_after = Pt(3)
                        p2.paragraph_format.line_spacing = 1.15
                        p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p2 = para.insert_paragraph_before(item_metode_penelitian)
                    p2.style.font.size = Pt(10)
                    p2.style.font.name = 'Arial'
                    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p2.paragraph_format.space_before = Pt(3)
                    p2.paragraph_format.space_after = Pt(3)
                    p2.paragraph_format.line_spacing = 1.15
                    p2.paragraph_format.first_line_indent = Inches(0.5)

        elif "AA5" in para.text:
            para.clear()
            for item_hasil in bagian.get('Hasil dan Pembahasan', ['Tidak ada hasil dan pembahasan ditemukan.']):
                if isinstance(item_hasil, list):
                    for i, list_item in enumerate(item_hasil, 1):
                        p3 = para.insert_paragraph_before(f"{i}. {list_item}")
                        p3.style.font.size = Pt(10)
                        p3.style.font.name = 'Arial'
                        p3.paragraph_format.left_indent = Inches(0.5)
                        p3.paragraph_format.first_line_indent = Inches(-0.15)
                        p3.paragraph_format.space_before = Pt(3)
                        p3.paragraph_format.space_after = Pt(3)
                        p3.paragraph_format.line_spacing = 1.15
                        p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p3 = para.insert_paragraph_before(item_hasil)
                    p3.style.font.size = Pt(10)
                    p3.style.font.name = 'Arial'
                    p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p3.paragraph_format.space_before = Pt(3)
                    p3.paragraph_format.space_after = Pt(3)
                    p3.paragraph_format.line_spacing = 1.15
                    p3.paragraph_format.first_line_indent = Inches(0.5)

        elif "AA6" in para.text:
            para.clear()
            for item_kesimpulan in bagian.get('Kesimpulan', ['Tidak ada kesimpulan ditemukan.']):
                if isinstance(item_kesimpulan, list):
                    for i, list_item in enumerate(item_kesimpulan, 1):
                        p4 = para.insert_paragraph_before(f"{i}. {list_item}")
                        p4.style.font.size = Pt(10)
                        p4.style.font.name = 'Arial'
                        p4.paragraph_format.left_indent = Inches(0.5)
                        p4.paragraph_format.first_line_indent = Inches(-0.15)
                        p4.paragraph_format.space_before = Pt(3)
                        p4.paragraph_format.space_after = Pt(3)
                        p4.paragraph_format.line_spacing = 1.15
                        p4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p5 = para.insert_paragraph_before(item_kesimpulan)
                    p5.style.font.size = Pt(10)
                    p5.style.font.name = 'Arial'
                    p5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p5.paragraph_format.space_before = Pt(3)
                    p5.paragraph_format.space_after = Pt(3)
                    p5.paragraph_format.line_spacing = 1.15
                    p5.paragraph_format.first_line_indent = Inches(0.5)

        elif "AA7" in para.text:
            para.clear()
            for item_referensi in bagian.get('Referensi', ['Tidak ada referensi ditemukan.']):
                if isinstance(item_referensi, list):
                    for i, list_item in enumerate(item_referensi, 1):
                        p6 = para.insert_paragraph_before(f"{i}. {list_item}")
                        p6.style.font.size = Pt(10)
                        p6.style.font.name = 'Arial'
                        p6.paragraph_format.left_indent = Inches(0.5) 
                        p6.paragraph_format.first_line_indent = Inches(-0.5)  
                        p6.paragraph_format.space_before = Pt(3)  
                        p6.paragraph_format.space_after = Pt(3)   
                        p6.paragraph_format.line_spacing = 1.15   
                        p6.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p6 = para.insert_paragraph_before(item_referensi)
                    p6.style.font.size = Pt(10)
                    p6.style.font.name = 'Arial'
                    p6.paragraph_format.left_indent = Inches(0.5) 
                    p6.paragraph_format.first_line_indent = Inches(-0.5)  
                    p6.paragraph_format.space_before = Pt(3)  
                    p6.paragraph_format.space_after = Pt(3)   
                    p6.paragraph_format.line_spacing = 1.15
                    p6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


        elif "Keyword: Maksimal 5 kata dari jurnal (dipisahkan dengan titik koma)" in para.text:
            para.clear()
            run = para.add_run(f"Keyword: {bagian.get('Kata Kunci', 'Tidak ada kata kunci ditemukan.')}")
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            run.italic = True
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        elif "Kata kunci: Maksimal 5 kata dari jurnal (dipisahkan dengan titik koma)" in para.text:
            para.clear()
            run = para.add_run(f"Kata kunci: {bagian.get('Kata Kunci', 'Tidak ada kata kunci ditemukan.')}")
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def simpan_dokumen_baru(dokumen_template, path_baru):
    dokumen_template.save(path_baru)

def konversi_skripsi_ke_jurnal(path_skripsi, path_template, path_output):
    dokumen_skripsi = Document(path_skripsi)
    dokumen_template = baca_template(path_template)
    bagian = ekstrak_bagian_dengan_nlp(dokumen_skripsi)
    sesuaikan_dengan_template(dokumen_template, bagian)
    simpan_dokumen_baru(dokumen_template, path_output)

path_skripsi = "prototype/data/skripsi_test_text-biasa.docx"
path_template = "prototype/template/Template.docx"
path_output = "prototype/output/skripsi_ke_jurnal.docx"

konversi_skripsi_ke_jurnal(path_skripsi, path_template, path_output)
print(f"Dokumen disimpan sebagai {path_output}")
